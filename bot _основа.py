import requests
from bs4 import BeautifulSoup
import re
import datetime
import dateutil.parser as dparser
from docx import Document
import os

def fillb(key, rows): #Заполняем пустые ячейки в бакалавриате 
    result = B.tables[-3].rows[rows].cells[1].text 
    result1 = B.tables[-3].rows[rows].cells[2].text 
    result2 = B.tables[-3].rows[rows].cells[3].text 
    student[key] = 'Сроки приема документов: '+result.replace('\n', ' ').replace('- ', '')+', Сроки проведения вступительных испытаний: '+result1.replace('\n', ' ').replace('- ', '')+', Примечания: '+result2.replace('\n', ' ').replace('- ', '') 
    return student[key] 

def fillm(key, rows): #Заполняем пустые ячейки в магистратуре 
    result = M.tables[-2].rows[rows].cells[1].text 
    result1 = M.tables[-2].rows[rows].cells[2].text 
    student[key] = 'Сроки приема документов: '+result.replace('\n', ' ').replace('- ', '')+', Сроки проведения вступительных испытаний: '+result1.replace('\n', ' ').replace('- ', '') 
    return student[key] 

def filla(key, rows): #Заполняем пустые ячейки в аспирантуре 
    result = A.tables[-3].rows[rows].cells[1].text 
    result1 = A.tables[-3].rows[rows].cells[2].text 
    student[key] = 'Сроки приема документов: '+result.replace('\n', ' ').replace('- ', '')+', Сроки проведения вступительных испытаний: '+result1.replace('\n', ' ').replace('- ', '') 
    return student[key]

def docx_reader(document): " #Красивый текст из файла docx
    text = ''
    for i in range (len(document.paragraphs)):
        text += document.paragraphs[i].text + '\n'
    return text.replace('\n'  + '\n', '\n')

def fragment_finder(text, begin_phrase, end_phrase): #Поиск фрагмента текста по началу и концу фрагмента (фразы)
    fragment_begin = text.find(begin_phrase)
    fragment_end = text.rfind(end_phrase)
    fragment = text[fragment_begin : fragment_end]
    return fragment.replace('\n', ' ')

def url_reader(url): #Достаём текст с сайта, используя его url
    page = requests.get(url)
    soup = BeautifulSoup(page.text, 'html.parser')
    text = soup.find(class_='text').text
    return text

#словарь из таблицы
student = dict(
    Б = ['Б_дог','Б_бюд'],
      Б_дог = ['Б_дог_оч','Б_дог_заоч','Б_дог_оч_заоч'],
        Б_дог_оч = ['Б_дог_оч_ДВИ','Б_дог_оч_неДВИ'],
          Б_дог_оч_ДВИ = [],
          Б_дог_оч_неДВИ = [],
        Б_дог_заоч = ['Б_дог_оч_заоч_ИДДО','Б_дог_оч_заоч_ИПЭЭф','Б_дог_оч_заоч_другие'],
          Б_дог_заоч_ИДДО = [],
          Б_дог_заоч_ИПЭЭф = [],
          Б_дог_заоч_другие = [],
        Б_дог_оч_заоч = ['Б_дог_оч_заоч_ИПЭЭф','Б_дог_оч_заоч_другие'],
          Б_дог_оч_заоч_ИПЭЭф = [],
          Б_дог_оч_заоч_другие = [],
      Б_бюд = ['Б_бюд_ДВИ','Б_бюд_неДВИ'],
        Б_бюд_ДВИ = [],
        Б_бюд_неДВИ = [],
    М = ['М_дог','М_бюд'],
      М_дог = ['М_дог_оч','М_дог_заоч','М_дог_оч_заоч'],
        М_дог_оч = [],
        М_дог_заоч = ['М_дог_заоч_ИДДО','М_дог_заоч_ИПЭЭф'],
          М_дог_заоч_ИДДО = [],
          М_дог_заоч_ИПЭЭф = [],
        М_дог_оч_заоч = ['М_дог_оч_заоч_ИнЭИ','М_дог_оч_заоч_ИПЭЭф_ИэТ_ГПИ'],
          М_дог_оч_заоч_ИнЭИ = [],
          М_дог_оч_заоч_ИПЭЭф_ИэТ_ГПИ = [],
      М_бюд = [],
    А = ['А_дог','А_бюд'],
      А_дог = [],
      А_бюд = ['А_бюд_обыч','А_бюд_доп'],
        А_бюд_обыч = [],
        А_бюд_доп = []
              )

B = Document('C:/Users/Vova/Desktop/dictionary_tables/Бакалавриат.docx')
M = Document('C:/Users/Vova/Desktop/dictionary_tables/Магистратура.docx')
A = Document('C:/Users/Vova/Desktop/dictionary_tables/Аспирантура.docx')

fillb('Б_дог_оч_ДВИ', 4)
fillb('Б_дог_оч_неДВИ', 5)
fillb('Б_дог_заоч_ИДДО', 8)
fillb('Б_дог_заоч_ИПЭЭф', 7)
fillb('Б_дог_заоч_другие', 9)
fillb('Б_дог_оч_заоч_ИПЭЭф', 7)
fillb('Б_дог_оч_заоч_другие', 6)
fillb('Б_бюд_ДВИ', 2)
fillb('Б_бюд_неДВИ', 3)
fillm('М_дог_оч', 2)
fillm('М_дог_заоч_ИДДО', 5)
fillm('М_дог_заоч_ИПЭЭф', 4)
fillm('М_дог_оч_заоч_ИнЭИ', 3)
fillm('М_дог_оч_заоч_ИПЭЭф_ИэТ_ГПИ', 4)
fillm('М_бюд', 2)
filla('А_дог', 2)
filla('А_бюд_обыч', 1)
filla('А_бюд_доп', 3)

B_text = docx_reader(B)
M_text = docx_reader(M)
A_text = docx_reader(A)

#Словари формата словарь[вопрос] = ответ на вопрос
    #Бакалавриат
B_questions = dict (
    Б_Что_писать_в_заявлении = fragment_finder(B_text, 'В заявлении о приеме', 'и месте их сдачи (если проведение подобных вступительных испытаний и место проведения предусмотрены)')
                               + '\n' + 
                               fragment_finder(B_text, 'Кроме указанных выше сведений в заявлении фиксируются', ' этого особого права только в МЭИ и только на одну образовательную программу'),
    Б_Куда_подавать = fragment_finder(B_text, 'Прием документов, необходимых для поступления, проводится в', 'Только после этого документы на поступление считаются принятыми.'),
    Б_Ограничение_количества_заявлений = fragment_finder(B_text, 'Поступающий на обучение по программам бакалавриата или программам специалитета имеет право подать заявления ', 'заявление о приеме на обучение (при одновременном поступлении в МЭИ и его филиалы - отдельно заявление в МЭИ (г. Москва) и в каждый филиал)'),
    Б_Какие_документы = fragment_finder(B_text, 'При подаче документов о приеме на обучение поступающий предъявляет', 'Республики Крым и города федерального значения Севастополя» и (или) Федеральным законом № 84-ФЗ.') 
                        + '\n' + 
                        fragment_finder(B_text, 'Вместе с заявлением поступающий сдает', 'документ иностранного государства об образовании')
                        + '\n' +
                        fragment_finder(B_text, 'Поступающий одновременно с подачей', '(будет подано) заявление о согласии на зачисление.')
                        + '\n' +
                        fragment_finder(B_text, 'Кроме документа об образовании поступающий', 'для лиц, поступающих в филиалы МЭИ по результатам вступительных испытаний, проводимых МЭИ самостоятельно')
                        + '\n' +
                        'Подробнее читайте на сайте Приёмной Комиссии: https://www.pkmpei.ru/info/documents.html' 
)                                           
    #Магистратура
M_questions = dict (
    М_Что_писать_в_заявлении = fragment_finder(M_text,'В заявлении о приеме на обучение поступающий указывает','и месте их сдачи (если проведение подобных вступительных испытаний и место проведения предусмотрены правилами проведения вступительных испытаний, проводимых МЭИ)')
                               + '\n' +
                               fragment_finder(M_text, 'Кроме указанных выше сведений в заявлении фиксируются с заверением ', 'зачисление, - обязательство представить соответствующие документы не позднее этого дня.'), 
    М_Куда_подавать = fragment_finder(M_text, 'Прием документов, необходимых для поступления, проводится в', 'Только после этого документы на поступление считаются принятыми.'),
    М_Какие_документы = fragment_finder(M_text, 'При подаче документов о приеме на обучение поступающий предъявляет', '') 
                        + '\n' +
                        fragment_finder(M_text, 'Кроме документа об образовании поступающий сдает в приемную комиссию копии (желательно при предъявлении оригиналов)', '4)	иные документы (предоставляются по усмотрению поступающего).')
                        + '\n' +
                        fragment_finder(M_text, 'Вместе с заявлением поступающий сдает оригинал или ксерокопию документа о высшем образовании', 'иностранном государстве, отвечающего требованиям')
                        + '\n' +
                        'Подробнее читайте на сайте Приёмной Комиссии: https://www.pkmpei.ru/info/documents_mag.html'
)    
    #Аспирантура
A_questions = dict (
    А_Что_писать_в_заявлении = fragment_finder(A_text, 'В Заявлении о приеме поступающий указывает ', '(в случае непоступления на обучение и в иных случаях, установленных Правилами приема).')
                               + '\n' +
                               fragment_finder(A_text, 'Кроме указанных выше сведений в Заявлении фиксируются', 'заверяются подписью поступающего (доверенного лица).'),
    A_Куда_подавать = fragment_finder(A_text, 'Прием документов, необходимых для поступления, проводится', 'Только после этого документы на поступление считаются поданными.'),
    A_Ограничения_количества_заявлений = fragment_finder(A_text, 'Поступающий вправе одновременно', 'различным условиям поступления поступающий подает одно Заявление о приеме'),
    A_Какие_документы = fragment_finder(A_text, 'При подаче документов, необходимых для поступления', '(заполняется на бланке МЭИ при подаче документов).')
                        + '\n' +
                        fragment_finder(A_text, 'При подаче Заявления о приеме кроме документов', 'оказании платных образовательных услуг с юридическим лицом).')
                        + '\n' +
                        fragment_finder(A_text, 'Поступающие могут ', 'подаваемых для поступления. Заверения копий указанных документов не требуется.')
)    

#Список формата "достижение - кол-во баллов", почти ответ на вопрос "Какие индивидуальные достижения дают дополнительные балллы и сколько?"
a = url_reader('https://www.pkmpei.ru/info/raitingrules')
a = a.split('\n\n')
individual_achievement  = []
for i in range(len(a)):  
  a[i] = a[i].replace('\xa0', '').replace('\r', '').split('\n')
for j in range(3, len(a) - 2):
  individual_achievement.append(a[j][2] + ' — ' + a[j][3])

class analyssaita:
    def __init__(self):
        self.page = requests.get('https://www.pkmpei.ru/')
        self.soup = BeautifulSoup(self.page.text, 'html.parser')
    def timetable(self):
        timetable1 = []
        timetable2 = []
        timetable3 = []
        timetable4 = []
        self.timetable = dict()
        s = ''
        l = 0
        timetable = str(self.soup.find(class_='title2'))
        timetable = re.split('; ', timetable)
        for i in timetable:
            if 'left' in i or 'center' in i:
                timetable1.append(i)
        for i in timetable1:
            if '<b>' in i:
                timetable2.append(i)
                timetable2 = re.split(';', str(timetable2))
                for i in timetable2:
                    if '"><b>' in i:
                        start = '"><b>'
                        end = '</b></td>\\n<td style="text-align: center'
                        i = i[i.find(start)+len(start):i.rfind(end)]
                        timetable3.append(i)
                for i in timetable2:
                    if '\\xa0' in i:
                        if 'Выходной' in i:
                            i = 'Выходной'
                        else:
                            i = re.findall('\d+', i)
                            for j in range(len(i) - 2):
                                if i[j] == '0':
                                    del i[j]
                                elif i[j] == '010':
                                    i[j] = '10'
                                elif i[-1] == '0':
                                    del i[-1]
                        i = ''.join(map(str, i))
                        timetable4.append(i)
                for i in range(len(timetable4)):
                    if timetable4[i] == '010001700':
                        timetable4[i] = '10:00 - 17:00'
                    elif timetable4[i] == '10001700':
                        timetable4[i] = '10:00 - 17:00'
                    elif timetable4[i] == '10001600':
                        timetable4[i] = '10:00 - 16:00'
        for i in range(len(timetable3)):
            self.timetable[l] = timetable3[i] + ': ' + timetable4[i]
            l = l +1
        print(self.timetable)
    def adress(self):
        adress = str(self.soup(class_="text"))
        start = 'Наш адрес'
        end = '/><br/><h1'
        adress = adress[adress.find(start)+len(start):adress.rfind(end)]
        start1 = '<div style="text-align: left;">'
        end1 = ', аудитория Б-209<br'
        adress1 = adress[adress.find(start1)+len(start1):adress.rfind(end1)]
        print(adress1)
        start2 = '<div style="text-align: left;">'
        end2 = '<br'
        self.adress2 = adress[adress.find(start2)+len(start2):adress.rfind(end2)]
        adress1 =adress1.replace(' ', '')
        self.adress1 = 'https://www.google.com/search?q=' + adress1
        print(self.adress1)
    def contactniedannie(self):
        contactniedannie = str(self.soup(class_="text"))
        start = 'Контактные данные'
        end = '</span><br/> </b></b>'
        start1 = '</h1> <br/> <b>'
        end1 = '<b><br/>'
        contactniedannie = contactniedannie[contactniedannie.find(start)+len(start):contactniedannie.rfind(end)] + 'JFKLDSJFLSD'
        self.telefon = contactniedannie[contactniedannie.find(start1)+len(start1):contactniedannie.rfind(end1)]
        print(self.telefon)
        start3 = '<span style="color: rgb(0, 0, 255);">'
        end3 = 'JFKLDSJFLSD'
        self.email = 'e-mail: ' + contactniedannie[contactniedannie.find(start3)+len(start3):contactniedannie.rfind(end3)]
        print(self.email)
    def novosti(self):
        novosti = str(self.soup.find_all(id="right_news_item"))
        novosti1 = []
        data = []
        data2 = []
        d = re.split(',', novosti)
        self.novosti = dict()
        urls = []
        news = []
        for i in d:
            if '<div id="right_news_item"' in i:
                novosti1.append(i)
       
        #print(novosti1)
        for i in range(len(novosti1)):
            start = 'href="'
            end ='">\n<div>'
            l = novosti1[i][novosti1[i].find(start)+len(start):novosti1[i].rfind(end)]
            url = 'https://www.pkmpei.ru/' + l
            urls.append(url)
        for i in urls:
            page = requests.get(i)
            soup = BeautifulSoup(page.text, 'html.parser')
            data1= str(soup.find(class_= 'title3'))
            start1 = '<div class="title3">'
            end1 = '</div>'
            data1 = data1[data1.find(start1)+len(start1):data1.rfind(end1)]
            data.append(data1)
            for i in data:
                if i not in data2:
                    data2.append(i)
            text = str(soup.find_all(class_= 'text'))
            start2 ='<h1>'
            end2 = '</h1>'
            text = text[text.find(start2)+len(start2):text.rfind(end2)]
            news.append(text)
            for i in range(len(news)):
                if '</h1>\n<br><h1>' in news[i]:
                    news[i] = news[i][news[i].find('<h1>'):]
                    news[i] = news[i][4:]
            for i in range(len(data2)):
                self.novosti[data2[i]] = 'НОВОСТИ:' + '\n' + data2[i] + '\n' + news[i] + '\n' + 'Ссылка: ' + urls[i]
            #print(self.novosti)
t = analyssaita()
print(t.timetable())
print(t.adress())
print(t.contactniedannie())
print(t.novosti())
import datetime
import random
import vk_api
from vk_api.longpoll import VkLongPoll, VkEventType
from vk_api.keyboard import VkKeyboard, VkKeyboardColor
with open('C:/Users/Vova/Desktop/secret.txt','r') as file:
    te=str(file.read())
vk_session = vk_api.VkApi(token=te)
vk = vk_session.get_api()
check=set()
text1 = ''
text2 = ''
text3 = ''
date=datetime.date.today()
def check_date(string,h):
    day=string[:2]
    try:
        day=int(day)
    except:
        day=int(day[1])
    month=string[3:5]
    try:
        month=int(month)
    except:
        month=int(month[1])
    year=int(string[6:10])
    days=day+month*30+year*365
    if date.day+date.month*30+date.year*365-days<=h:
        return True
    else:
        return False
def send_to_all(text):
    members=vk.groups.getMembers(group_id='183149235')['items']
    for person in members:
        try:
            if person not in ignor:
                vk.messages.send(user_id=person,
                        random_id=random.randint(1,10**9),
                        message=text)
        except:
            continue
def create_keyboard():
   
    keyboard = VkKeyboard(one_time=False)

    keyboard.add_button('Контактные данные', color=VkKeyboardColor.POSITIVE)
    keyboard.add_button('Отписаться/подписаться на новости',color=VkKeyboardColor.POSITIVE)
    keyboard.add_line()  
    keyboard.add_button('График работы приемной комиссии', color=VkKeyboardColor.NEGATIVE)
 
    keyboard.add_line()
    keyboard.add_button('Адрес', color=VkKeyboardColor.PRIMARY)
    keyboard.add_button('Новости', color=VkKeyboardColor.PRIMARY)

    keyboard.add_line()
    keyboard.add_button('Информация о поступлении',color=VkKeyboardColor.NEGATIVE)
    keyboard = keyboard.get_keyboard()

    return keyboard
def create_klav(knopki):
    keyboard = VkKeyboard(one_time=True)
    for b in knopki:
        if b=='new_line':
            keyboard.add_line()
        else:
            keyboard.add_button(b, color=VkKeyboardColor.POSITIVE)
    keyboard.add_line()
    keyboard.add_button('На главную',color=VkKeyboardColor.NEGATIVE)
    keyboard = keyboard.get_keyboard()
    return keyboard
def create_grafic():
    keyboard = VkKeyboard(one_time=True)
    keyboard.add_button('На сегодня',color=VkKeyboardColor.PRIMARY)
    keyboard.add_button('На эту неделю',color=VkKeyboardColor.PRIMARY)
    keyboard=keyboard.get_keyboard()
    return keyboard
flag=True
date=datetime.datetime.today()
day=date.weekday()
print(day)
ignor=set()
while True:
    try:
        for event in VkLongPoll(vk_session).listen():
            keyboard = create_keyboard()
            keyboard1 = create_grafic()
            date=datetime.datetime.today()
            day=date.weekday()
            #print(day)
            date=datetime.date.today()
            if day==6 and flag==True:
                flag=False
                nov=''
                for key in t.novosti:
                    if check_date(key,180):
                        nov+=t.novosti[key]+'\n'+'\n'
                try:
                    send_to_all(nov)
                except:
                    print('Нет новостей')
            elif day!=6:
                flag=True
            if event.type == VkEventType.MESSAGE_NEW and event.to_me and event.text:
                if event.from_user:
                    text=event.text
                    text=text.lower()
                    if text=='работает ли сегодня приемная комиссия?' or text =='на сегодня':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=t.timetable[day],
                                         keyboard=keyboard)
                    elif text =='когда вообще работает приемная комиссия?' or text =='на эту неделю':
                        for i in range(len(t.timetable)):
                            text1 = text1 + t.timetable[i] + '\n'
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=text1, keyboard = keyboard)
                        text1 = ''
                    elif text == 'адрес' or text == 'какой адрес?' or text == 'скажи адрес':
                        text3 = 'Адрес: ' + t.adress2 + '\n' + 'Как дойти: ' + t.adress1
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=text3, keyboard = keyboard)
                    elif text == 'контактная информация' or text =='контактные данные':
                        text2 = t.telefon + '\n' + t.email
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=text2, keyboard = keyboard)
                    elif text == 'начать':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Здравствуйте! Я очень рад Вам!', keyboard = keyboard)
                    elif text == 'новости':
                        nov=''
                        for key in t.novosti:
                            #print(key)
                            if check_date(key,180):
                                nov+=t.novosti[key]+'\n'+'\n'
                        
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=nov,
                                         keyboard = keyboard)
                    elif text == 'информация о поступлении':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Какую ступень образования Вы хотите получить?',
                                         keyboard = create_klav(['Бакалавриат','Магистратура','Аспирантура']))
                    elif text == 'график работы приемной комиссии':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Выберите: ',
                                         keyboard = keyboard1)
                    elif text == 'отписаться/подписаться на новости':
                        if event.user_id not in ignor:
                            ignor.add(event.user_id)
                            vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Теперь Вы отписаны от автоматичечкой рассылки новостей',
                                         keyboard = keyboard)

                        else:
                            ignor.remove(event.user_id)
                            vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Теперь Вы подписаны на автоматичечкую рассылку новостей',
                                         keyboard = keyboard)
                    elif text == 'аспирантура':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете на бюджет или по договору?', keyboard = create_klav(['Бюджет(А)','По договору(А)']))
                    elif text =='бюджет(а)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете по целевому набору или по дополнительному?', keyboard = create_klav(['Целевой набор(А)','Дополнительный набор(А)']))
                    elif text =='целевой набор(а)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['А_бюд_обыч'], keyboard = keyboard)
                    elif text =='дополнительный набор(а)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['А_бюд_доп'], keyboard = keyboard)
                    elif text =='по договору(а)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['А_дог'], keyboard = keyboard)
                    elif text =='магистратура':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете на бюджет или по договору?', keyboard = create_klav(['Бюджет(М)','По договору(М)']))
                    elif text == 'бюджет(м)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_бюд'], keyboard = keyboard)
                    elif text == 'по договору(м)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Какую форму обучения вы выбираете?', keyboard = create_klav(['Очная(М)','Заочная(М)','Очно-заочная(М)']))
                    elif text == 'очная(м)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_дог_оч'], keyboard = keyboard)
                    elif text =='заочная(м)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете в ИДДО или в ИПЭЭф ?', keyboard = create_klav(['ИДДО(М_з)','ИПЭЭф(М_з)']))
                    elif text == 'иддо(м_з)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_дог_заоч_ИДДО'], keyboard = keyboard)
                    elif text == 'ипээф(м_з)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_дог_заоч_ИПЭЭф'], keyboard = keyboard)
                    elif text == 'очно-заочная(м)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете в ИнЭИ, ИПЭЭф, ИЭТ или ГПИ?', keyboard = create_klav(['ИнЭИ(М_оз)','ИПЭЭф, ИЭТ или ГПИ(М_оз)']))
                    elif text == 'инэи(м_оз)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_дог_оч_заоч_ИнЭИ'], keyboard = keyboard)
                    elif text == 'ипээф, иэт или гпи(м_оз)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['М_дог_оч_заоч_ИПЭЭф_ИэТ_ГПИ'], keyboard = keyboard)
                    elif text == 'бакалавриат':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете на бюджет или по договору?', keyboard = create_klav(['Бюджет(Б)','По договору(Б)'])) 
                    elif text =='бюджет(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете со вступительными экзаменами или без?', keyboard = create_klav(['Со вступительными экзаменами(б)','Без вступительных экзаменов(б)']))
                    elif text =='со вступительными экзаменами(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_бюд_ДВИ'], keyboard = keyboard)
                    elif text =='без вступительных экзаменов(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_бюд_неДВИ'], keyboard = keyboard)
                    elif text =='по договору(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Какую форму обучения Вы выбираете?', keyboard = create_klav(['Очная(Б)','Заочная(Б)','Очно-заочная(Б)']))
                    elif text == 'очная(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете со вступительными экзаменами или без?', keyboard = create_klav(['Со вступительными экзаменами(б_д_о)','Без вступительных экзаменов(б_д_о)']))
                    elif text == 'со вступительными экзаменами(б_д_о)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_оч_ДВИ'], keyboard = keyboard)
                    elif text == 'без вступительных экзаменов(б_д_о)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_оч_неДВИ'], keyboard = keyboard)
                    elif text == 'заочная(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете в ИДДО, ИПЭЭф или в другой институт? ', keyboard = create_klav(['ИПЭЭф(б_д_з)','ИДДО(б_д_з)','другой институт(б_д_з)']))
                    elif text == 'ипээф(б_д_з)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_заоч_ИПЭЭф'], keyboard = keyboard)
                    elif text == 'иддо(б_д_з)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_заоч_ИДДО'], keyboard = keyboard)
                    elif text == 'другой институт(б_д_з)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_заоч_другие'], keyboard = keyboard)
                    elif text == 'очно-заочная(б)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы поступаете в ИПЭЭф или в другой институт?', keyboard = create_klav(['ИПЭЭФ(б_д_оз)','другой институт(б_д_оз)']))
                    elif text == 'ипээф(б_д_оз)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_оч_заоч_ИПЭЭф'], keyboard = keyboard)
                    elif text == 'другой институт(б_д_оз)':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message=student['Б_дог_оч_заоч_другие'], keyboard = keyboard)
                    elif text == 'на главную':
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Вы перемещены на главную страницу', keyboard = keyboard)
                    else:
                        vk.messages.send(user_id=event.user_id,
                                         random_id=random.randint(1,10**9),
                                         message='Извините, я вас не понимаю. Используйте команды, которые вызываются с помощью кнопок снизу.', keyboard = keyboard)
    except:
        if day==6 and flag==True:
            flag=False
            nov=''
            for key in t.novosti:
                if check_date(key,180):
                    nov+=t.novosti[key]+'\n'+'\n'
            try:
                send_to_all(nov)
            except:
                print('Нет новостей')
        elif day!=6:
            flag=True
        continue
